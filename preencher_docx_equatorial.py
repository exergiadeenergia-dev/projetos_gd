#!/usr/bin/env python3
"""
Preenche os campos de IDENTIFICAÇÃO (cliente, UC, endereço, coordenadas,
responsável técnico, equipamentos) do Memorial Descritivo e do Relatório de
Comissionamento da Equatorial-GO, usando um modelo de referência genérico
de referência (substituição de texto).

IMPORTANTE — o que este script NÃO faz (fica por conta do engenheiro):
As seções de dimensionamento técnico (disjuntor CA/CC, DPS, aterramento,
bitola de cabos, levantamento de carga) dependem de cálculo de engenharia
específico de cada projeto e não são preenchidas automaticamente. Essas
seções continuam com os valores do projeto de referência (modelo genérico) e
precisam ser revisadas manualmente para cada novo cliente.

Uso:
    python3 preencher_docx_equatorial.py dados_cliente.json memorial_referencia.docx saida_memorial.docx
    python3 preencher_docx_equatorial.py dados_cliente.json comissionamento_referencia.docx saida_comissionamento.docx --tipo comissionamento
"""
import sys
import json
from docx import Document


def substituir_foto_localizacao(doc, caminho_imagem):
    """Troca a foto de satélite/localização do imóvel (a que aparece logo
    após 'Coordenadas georrefenciadas' e antes de 'Figura 1: Localização da
    unidade consumidora'), mantendo o restante do documento intacto."""
    for p in doc.paragraphs:
        for run in p.runs:
            blips = run._element.findall(
                ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
            )
            for b in blips:
                rid = b.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                part = doc.part.related_parts.get(rid)
                if part is not None and part.content_type == "image/png" and len(part.blob) > 500_000:
                    with open(caminho_imagem, "rb") as f:
                        part._blob = f.read()
                    return True
    return False


def replace_everywhere(doc, replacements):
    """Substitui strings em parágrafos e células de tabela, preservando
    formatação ao concatenar todos os runs de cada parágrafo antes de
    substituir (evita problemas de texto quebrado entre runs)."""

    def merge_split_paragraphs(paragraphs):
        """Alguns valores de referência (ex: nome do modelo) aparecem
        quebrados entre dois parágrafos consecutivos no documento original
        (herança de formatação colorida). Se um valor de referência só é
        encontrado ao juntar dois parágrafos vizinhos, funde o texto do
        segundo no primeiro antes de substituir."""
        texts = [p.text for p in paragraphs]
        for old in replacements:
            if not old or len(old) < 4:
                continue
            for i in range(len(paragraphs) - 1):
                if old in texts[i]:
                    continue
                for sep in ("", "\n"):
                    combined = texts[i] + sep + texts[i + 1]
                    if old in combined and old not in texts[i + 1]:
                        p = paragraphs[i]
                        if p.runs:
                            p.runs[0].text = combined
                            for run in p.runs[1:]:
                                run.text = ""
                        for run in paragraphs[i + 1].runs:
                            run.text = ""
                        texts[i] = combined
                        texts[i + 1] = ""
                        break

    def process_paragraph(p):
        full_text = "".join(run.text for run in p.runs)
        if not full_text:
            return
        new_text = full_text
        for old, new in replacements.items():
            if old and old in new_text:
                new_text = new_text.replace(old, str(new))
        if new_text != full_text:
            if p.runs:
                p.runs[0].text = new_text
                for run in p.runs[1:]:
                    run.text = ""

    merge_split_paragraphs(doc.paragraphs)
    for p in doc.paragraphs:
        process_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                merge_split_paragraphs(cell.paragraphs)
                for p in cell.paragraphs:
                    process_paragraph(p)


def fix_tabela_gerador(doc, quantidade_modulos, potencia_kw):
    """Corrige diretamente (por índice de célula, não por busca de texto —
    números curtos como '8' são perigosos demais para substituição global)
    as linhas 'Quantidade' e 'Potência do gerador [kW]' da Tabela 3."""
    t = doc.tables[0]
    t.rows[12].cells[1].paragraphs[0].runs[0].text = str(quantidade_modulos)
    for run in t.rows[12].cells[1].paragraphs[0].runs[1:]:
        run.text = ""
    pot_str = f"{potencia_kw:.2f}".replace(".", ",")
    t.rows[13].cells[1].paragraphs[0].runs[0].text = pot_str
    for run in t.rows[13].cells[1].paragraphs[0].runs[1:]:
        run.text = ""


def montar_substituicoes_memorial(dados):
    c = dados["cliente"]
    ref = dados["_referencia_hakknner"]
    m = c["modulos"][0]
    pot_kw = sum(mm["potencia_w"] * mm["quantidade"] for mm in c["modulos"]) / 1000
    equip_desc = " + ".join(
        f'{mm["quantidade"]} módulos {mm["fabricante"]}/{mm["modelo"]}' for mm in c["modulos"]
    )

    inv = c["inversores"][0]
    return {
        ref["nome"]: c["nome"],
        ref["conta_contrato"]: c["conta_contrato"],
        ref["endereco_uc"]: c["endereco"],
        ref["bairro_cidade"]: f'{c["bairro"]}. {c["municipio"]} – {c["uf"]}.',
        ref["cep"]: str(c["cep"]),
        ref["cidade"]: c["municipio"],
        ref["coordenadas"]: f'UTM {c["fuso"]}K, X: {c["coord_x"]} Y: {c["coord_y"]}',
        ref["equip_titulo"]: f'{pot_kw:g} kW',
        ref["equip_desc"]: equip_desc,
        ref["modalidade"]: c["modalidade_compensacao"],
        ref["modalidade"].upper(): c["modalidade_compensacao"].upper(),
        ref["mes_ano"]: c["mes_ano_documento"],
        ref["estado"]: c["uf_extenso"],
        f'Fabricante {ref["fab_modulo"]}': f'Fabricante {m["fabricante"]}',
        ref["fab_modulo"]: m["fabricante"],
        ref["modelo_modulo"]: m["modelo"],
        ref["fab_inversor_typo_memorial"]: inv["fabricante"],
        ref["modelo_inversor"]: inv["modelo"],
    }


def montar_substituicoes_comissionamento(dados):
    c = dados["cliente"]
    ref = dados["_referencia_hakknner"]
    m = c["modulos"][0]
    inv = c["inversores"][0]

    return {
        ref["nome_maiusculo"]: c["nome"].upper(),
        ref["telefone_cliente"]: c["telefone_cliente"],
        ref["uc_pontuado"]: c["conta_contrato"],
        ref["data_conclusao"]: c["data_conclusao_geradora"],
        ref["endereco_completo"]: c["endereco_completo_comissionamento"],
        ref["cep"]: str(c["cep"]),
        ref["cidade"]: c["municipio"],
        ref["fab_modulo"]: m["fabricante"],
        ref["modelo_modulo"]: m["modelo"],
        ref["fab_inversor"]: inv["fabricante"],
        ref["modelo_inversor"]: inv["modelo"],
    }


if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Uso: python3 preencher_docx_equatorial.py dados.json referencia.docx saida.docx [--tipo comissionamento]")
        sys.exit(1)
    json_path, ref_path, saida_path = sys.argv[1:4]
    tipo = "memorial"
    if "--tipo" in sys.argv:
        tipo = sys.argv[sys.argv.index("--tipo") + 1]

    with open(json_path, encoding="utf-8") as f:
        dados = json.load(f)

    doc = Document(ref_path)
    if tipo == "comissionamento":
        subs = montar_substituicoes_comissionamento(dados)
    else:
        subs = montar_substituicoes_memorial(dados)

    replace_everywhere(doc, subs)
    fix_tabela_gerador(
        doc,
        sum(m["quantidade"] for m in dados["cliente"]["modulos"]),
        sum(m["potencia_w"] * m["quantidade"] for m in dados["cliente"]["modulos"]) / 1000,
    )

    if "--foto" in sys.argv:
        foto_path = sys.argv[sys.argv.index("--foto") + 1]
        ok = substituir_foto_localizacao(doc, foto_path)
        print(f"Foto de localização {'substituída' if ok else 'NÃO encontrada para substituir'}.")

    doc.save(saida_path)
    print(f"Documento salvo em: {saida_path}")
    print("\n⚠️  Lembrete: as seções de dimensionamento (disjuntor CA/CC, DPS, ")
    print("aterramento, bitola de cabos, levantamento de carga) NÃO foram")
    print("recalculadas — revise manualmente antes de enviar.")
